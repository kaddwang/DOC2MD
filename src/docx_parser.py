"""
Docx Parser Module
──────────────────
Reads a .docx file and extracts text paragraphs, tables, and inline images
in their original document order.  Each element is returned as a dictionary:

    {"type": "text",  "content": "paragraph text ..."}
    {"type": "image", "content": "<base64 string>", "mime": "image/png"}
    {"type": "text",  "content": "| col1 | col2 |\\n|---|---|\\n..."}  (table as Markdown)
"""

from __future__ import annotations

import base64
import re
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table


# ── Helpers ──────────────────────────────────────────────────

def _get_image_base64(image_part) -> Tuple[str, str]:
    """Extract base64-encoded data and MIME type from an image part."""
    blob: bytes = image_part.blob
    content_type: str = image_part.content_type
    b64_str = base64.b64encode(blob).decode("utf-8")
    return b64_str, content_type


def _dedup_text(text: str) -> str:
    """
    Fix duplicated text that some .docx files produce.

    Some Word documents contain paragraphs where the visible text appears
    twice back-to-back (e.g. "Release notesRelease notes").  This function
    detects and removes the duplication while preserving intentionally
    repeated content.
    """
    lines = text.split("\n")
    cleaned: List[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append(line)
            continue

        # Check if the line is exactly a doubled string
        length = len(stripped)
        if length >= 4 and length % 2 == 0:
            half = length // 2
            first_half = stripped[:half]
            second_half = stripped[half:]
            if first_half == second_half:
                cleaned.append(line.replace(stripped, first_half))
                continue

        # Check for partial duplication patterns (sentence-level)
        # e.g. "Some long sentence here.Some long sentence here."
        deduped = _dedup_sentence(stripped)
        if deduped != stripped:
            cleaned.append(line.replace(stripped, deduped))
        else:
            cleaned.append(line)

    return "\n".join(cleaned)


def _dedup_sentence(text: str) -> str:
    """
    Remove sentence-level duplication.
    Tries progressively shorter prefix lengths to find duplications.
    """
    length = len(text)
    if length < 6:
        return text

    # Try to find a split point where first half == second half
    # Allow for minor length differences (±2 chars)
    for half in range(length // 2 + 2, max(2, length // 2 - 2), -1):
        if half >= length:
            continue
        candidate = text[:half]
        remainder = text[half:]
        # Check if remainder starts with (or equals) the candidate
        if remainder == candidate or remainder.rstrip() == candidate.rstrip():
            return candidate

    return text


def _extract_paragraph_elements(paragraph, rels) -> List[Dict[str, str]]:
    """
    Walk through the XML of a single paragraph and yield text / image
    elements *in the order they appear*.
    """
    elements: List[Dict[str, str]] = []
    text_buffer: List[str] = []

    for child in paragraph._element:
        if child.tag == qn("w:r"):
            drawings = child.findall(f".//{qn('a:blip')}")
            if drawings:
                if text_buffer:
                    joined = "".join(text_buffer).strip()
                    if joined:
                        elements.append({"type": "text", "content": _dedup_text(joined)})
                    text_buffer.clear()

                for blip in drawings:
                    embed_id = blip.get(qn("r:embed"))
                    if embed_id and embed_id in rels:
                        rel = rels[embed_id]
                        image_part = rel.target_part
                        b64, mime = _get_image_base64(image_part)
                        elements.append(
                            {"type": "image", "content": b64, "mime": mime}
                        )
            else:
                # Only use <w:t> children — avoid doubling with child.text
                run_texts: List[str] = []
                for t_elem in child.findall(qn("w:t")):
                    if t_elem.text:
                        run_texts.append(t_elem.text)
                if run_texts:
                    text_buffer.append("".join(run_texts))

    if text_buffer:
        joined = "".join(text_buffer).strip()
        if joined:
            elements.append({"type": "text", "content": _dedup_text(joined)})

    return elements


def _table_to_markdown(table: Table) -> str:
    """
    Convert a python-docx Table object into a Markdown table string.
    """
    rows_data: List[List[str]] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = cell.text.strip().replace("\n", " <br> ")
            cells.append(text)
        rows_data.append(cells)

    if not rows_data:
        return ""

    lines: List[str] = []
    header = "| " + " | ".join(rows_data[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows_data[0]) + " |"
    lines.append(header)
    lines.append(separator)
    for row in rows_data[1:]:
        while len(row) < len(rows_data[0]):
            row.append("")
        line = "| " + " | ".join(row) + " |"
        lines.append(line)

    return "\n".join(lines)


def _iter_block_items(document):
    """
    Yield each paragraph and table in *document order*.
    """
    body = document.element.body
    for child in body:
        if child.tag == qn("w:p"):
            yield ("paragraph", child)
        elif child.tag == qn("w:tbl"):
            yield ("table", child)


# ── Public API ───────────────────────────────────────────────

def parse_docx(file_path: str | Path) -> List[Dict[str, str]]:
    """
    Parse a .docx file and return an ordered list of content elements,
    including paragraphs, inline images, and tables — all in document order.

    Parameters
    ----------
    file_path : str | Path
        Path to the .docx file.

    Returns
    -------
    list[dict]
        Each dict has ``type`` ("text" | "image") and ``content``.
        Image dicts additionally carry a ``mime`` key.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if file_path.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {file_path}")

    doc = Document(str(file_path))
    rels = doc.part.rels

    para_map = {}
    for p in doc.paragraphs:
        para_map[id(p._element)] = p

    table_map = {}
    for t in doc.tables:
        table_map[id(t._element)] = t

    all_elements: List[Dict[str, str]] = []

    for block_type, block_xml in _iter_block_items(doc):
        if block_type == "paragraph":
            para_obj = para_map.get(id(block_xml))
            if para_obj:
                para_elements = _extract_paragraph_elements(para_obj, rels)
                all_elements.extend(para_elements)

        elif block_type == "table":
            table_obj = table_map.get(id(block_xml))
            if table_obj:
                md_table = _table_to_markdown(table_obj)
                if md_table.strip():
                    all_elements.append({"type": "text", "content": md_table})

    # Merge consecutive text elements for cleaner payloads
    merged: List[Dict[str, str]] = []
    for elem in all_elements:
        if (
            elem["type"] == "text"
            and merged
            and merged[-1]["type"] == "text"
        ):
            merged[-1]["content"] += "\n\n" + elem["content"]
        else:
            merged.append(elem)

    return merged
